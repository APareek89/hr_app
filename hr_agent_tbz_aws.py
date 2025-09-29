#!/usr/bin/env python
# coding: utf-8

# In[ ]:


import os
import webbrowser
import re
import yaml
from flask import Flask, request, render_template_string, jsonify
from docx import Document
import pdfplumber
import threading
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import glob
import logging
import sys
import csv
import openai
import subprocess
from tempfile import NamedTemporaryFile
import time
import math
import json

# ====== CONFIGURE YOUR API KEY HERE ======
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY', '')
# =========================================

def get_base_dir():
    if getattr(sys, 'frozen', False):
        return getattr(sys, '_MEIPASS', os.getcwd())
    try:
        return os.path.dirname(os.path.abspath(__file__))
    except NameError:
        return os.getcwd()

BASE_DIR = get_base_dir()

log_file = os.path.join(BASE_DIR, 'app.log')
logging.basicConfig(level=logging.DEBUG,
                    format='%(asctime)s - %(levelname)s - %(message)s',
                    handlers=[logging.FileHandler(log_file), logging.StreamHandler()])
logger = logging.getLogger(__name__)

app = Flask(__name__)

def load_agent_config():
    default_config = {
        'models': {
            'embedding_model': 'text-embedding-3-small',
            'llm_model': 'gpt-4o-mini'
        },
        'weights': {
            # adjust weights for final score (must sum to 1.0)
            'job_features': 0.5,
            'industry': 0.15,
            'role': 0.15,
            'years': 0.20
        },
        'prompts': {
            'cleaning': 'Remove extra whitespaces, newlines, and non-alphanumeric noise. Keep core content like skills, experience, education.',
            'refinement': """You are a recruitment agent. Given a job description: "{job_desc}"
And a candidate resume summary: "{resume_summary}"
Mention in bullet points with each in separate line and headers: Candidate Name, Years of Experience, Contact Number, Contact Email, Current Company, Current Designation, and a short paragraph of 2-3 lines explaining why this candidate is a good fit, highlighting key matches in skills, experience, and qualifications.""",
            'job_desc_refined': """Extract key skill words from job responsibilities, activities, tasks to be performed and key activities provided in job description. Output only the bullets or short comma-separated keywords (no extra text).""",
            'resume_exp_extract': """Extract key skill words from job responsibilities, activities, tasks to be performed and key activities from this resume text. Output only the bullets or short comma-separated keywords (no extra text)."""
        }
    }
    config_path = os.path.join(BASE_DIR, 'configs', 'agent_config.yml')
    try:
        with open(config_path, 'r') as f:
            return yaml.safe_load(f)
    except Exception:
        logger.warning("Using default agent config")
        return default_config

agent_config = load_agent_config()

# Initialize OpenAI client if key provided
if OPENAI_API_KEY and OPENAI_API_KEY != "YOUR_API_KEY_GOES_HERE":
    os.environ['OPENAI_API_KEY'] = OPENAI_API_KEY
    openai_client = openai.OpenAI()
    logger.info("OpenAI client initialized")
else:
    openai_client = None
    logger.info("OpenAI key not configured; LLM/embedding calls will use placeholders.")

def convert_doc_to_docx(doc_path):
    out_dir = os.path.dirname(doc_path)
    try:
        subprocess.run(
            ['soffice', '--headless', '--convert-to', 'docx', '--outdir', out_dir, doc_path],
            check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
        )
        converted = os.path.splitext(doc_path)[0] + '.docx'
        return converted if os.path.exists(converted) else None
    except Exception as e:
        logger.error(f"Conversion failed for {doc_path}: {e}")
        return None

def clean_text(text):
    text = re.sub(r'\s+', ' ', text or '')
    text = re.sub(r'\[Page \d+\]', '', text)
    return text.strip() or ""

def extract_pdf_text(file_path_or_stream):
    try:
        with pdfplumber.open(file_path_or_stream) as pdf:
            text = ' '.join([page.extract_text() or '' for page in pdf.pages])
        return clean_text(text)
    except Exception as e:
        logger.error(f"PDF extract error: {e}")
        return ""

def get_embedding(text, model):
    text = clean_text(text)
    if not text:
        return np.zeros(1536) if openai_client else np.zeros(512)
    if openai_client is None:
        # deterministic placeholder embedding (not ideal but usable)
        arr = np.frombuffer(text.encode('utf-8')[:512].ljust(512, b'\0'), dtype=np.uint8).astype(float)
        norm = np.linalg.norm(arr)
        return arr / (norm if norm else 1.0)
    try:
        resp = openai_client.embeddings.create(model=model, input=text)
        return np.array(resp.data[0].embedding)
    except Exception as e:
        logger.error(f"Embedding error: {e}")
        raise

def llm_call(prompt, model):
    if openai_client is None:
        # fallback: return first 300 chars as simple summary
        return prompt[:300]
    try:
        resp = openai_client.chat.completions.create(
            model=model,
            messages=[{"role":"system","content":"You are a helpful recruitment assistant."},
                      {"role":"user","content":prompt}],
            max_tokens=512
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        logger.exception("LLM call failed")
        return ""

def extract_text_from_uploaded_file(file_storage):
    tmp_path = None
    converted = None
    try:
        fname = file_storage.filename or "upload"
        lower = fname.lower()
        if lower.endswith('.docx'):
            tmp = NamedTemporaryFile(delete=False, suffix='.docx')
            file_storage.save(tmp.name)
            tmp.close()
            doc = Document(tmp.name)
            text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            tmp_path = tmp.name
            return clean_text(text)
        elif lower.endswith('.doc'):
            tmp = NamedTemporaryFile(delete=False, suffix='.doc')
            file_storage.save(tmp.name)
            tmp.close()
            tmp_path = tmp.name
            converted = convert_doc_to_docx(tmp_path)
            if converted:
                doc = Document(converted)
                text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                return clean_text(text)
            else:
                return ""
        elif lower.endswith('.pdf'):
            tmp = NamedTemporaryFile(delete=False, suffix='.pdf')
            file_storage.save(tmp.name)
            tmp.close()
            text = extract_pdf_text(tmp.name)
            tmp_path = tmp.name
            return clean_text(text)
        else:
            return ""
    finally:
        try:
            if tmp_path and os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass
        try:
            if converted and os.path.exists(converted):
                os.remove(converted)
        except Exception:
            pass

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/process', methods=['POST'])
def process():
    try:
        # Collect inputs
        jd_text = request.form.get('text_input', '').strip()
        jd_file = request.files.get('jd_file')  # optional uploaded job description file
        years_exp_input = request.form.get('years_exp', '').strip()  # e.g., "5" or "3-6"
        industry_input = request.form.get('industry', '').strip()
        role_name = request.form.get('role_name', '').strip()

        # If jd_file provided, extract text and prefer it over free text
        if jd_file and jd_file.filename:
            file_text = extract_text_from_uploaded_file(jd_file)
            if file_text:
                jd_text = file_text if jd_text == "" else jd_text + "\n\n" + file_text

        if not jd_text:
            return jsonify({'output': 'Please provide a job description (text or upload).'}), 400

        # Build job features via LLM: responsibilities/activities/skills bullets
        job_prompt = agent_config['prompts']['job_desc_refined'].strip() + "\n\nJob Description:\n" + jd_text+agent_config['prompts']['refinement'].strip()
        #Promt to extract key words from job description
        job_features_prompt = agent_config['prompts']['job_desc_refined'].strip() + "\n\nJob Description:\n" + jd_text
        #Promt to extract key words from candidates profiles for similarity analysis
        job_experience_prompt = agent_config['prompts']['resume_exp_extract'].strip()

        job_features_text = llm_call(job_features_prompt, agent_config['models']['llm_model'])
        job_features_text = clean_text(job_features_text)
        logger.info(f"Job features (extracted): {job_features_text[:500]}")

        #job_experience_text = llm_call(job_experience_prompt, agent_config['models']['llm_model'])
        #job_experience_text = clean_text(job_experience_text)
        #logger.info(f"Job features (extracted): {job_features_text[:500]}")

        # Create combined job vector: embed job_features_text + industry + role_name
        embedding_model = agent_config['models']['embedding_model']
        job_feat_emb = get_embedding(job_features_text, embedding_model)
        #job_exp_emb = get_embedding(job_experience_text, embedding_model)
        industry_emb = get_embedding(industry_input or "", embedding_model)
        role_emb = get_embedding(role_name or "", embedding_model)
        #years_emb = get_embedding(years_exp_input or "", embedding_model)

        # Normalize years input into numeric (midpoint if range)
        def parse_years(s):
            s = s.strip()
            if not s: return None
            m = re.match(r'^\s*(\d+)\s*-\s*(\d+)\s*$', s)
            if m:
                a = int(m.group(1)); b = int(m.group(2))
                return (a + b) / 2.0
            m2 = re.match(r'^\s*(\d+)\s*$', s)
            if m2:
                return float(m2.group(1))
            return None
        target_years = parse_years(years_exp_input)

        # Load resumes from uploads + resumes/ folder, extract text and LLM-extracted experience bullets
        uploaded_files = request.files.getlist('file_input')
        resume_folder = os.path.join(BASE_DIR, 'resumes')
        os.makedirs(resume_folder, exist_ok=True)
        folder_files = glob.glob(os.path.join(resume_folder, '*.docx')) + glob.glob(os.path.join(resume_folder, '*.pdf')) + glob.glob(os.path.join(resume_folder, '*.doc'))

        if not uploaded_files and not folder_files:
            return jsonify({'output': 'No resumes uploaded and none found in resumes/ folder.'}), 400

        # helper to load resume text from path
        def load_resume_text_from_path(path):
            pl = path.lower()
            try:
                if pl.endswith('.docx'):
                    doc = Document(path)
                    return clean_text('\n'.join([p.text for p in doc.paragraphs if p.text.strip()]))
                elif pl.endswith('.doc'):
                    converted = convert_doc_to_docx(path)
                    if converted:
                        doc = Document(converted)
                        text = clean_text('\n'.join([p.text for p in doc.paragraphs if p.text.strip()]))
                        try:
                            os.remove(converted)
                        except Exception:
                            pass
                        return text
                    else:
                        return ""
                elif pl.endswith('.pdf'):
                    return extract_pdf_text(path)
            except Exception as e:
                logger.exception(f"Failed to read resume {path}: {e}")
            return ""

        resumes = {}  # fname -> {text, resume_features_text, embeddings... , years (if extracted)}
        # process uploaded files
        for f in uploaded_files:
            if not f or not (f.filename or "").strip():
                continue
            fname = f.filename
            lower = fname.lower()
            if not lower.endswith(('.doc', '.docx', '.pdf')):
                logger.info(f"Skipping unsupported upload: {fname}")
                continue
            text = extract_text_from_uploaded_file(f)
            if not text:
                logger.warning(f"No text extracted from uploaded resume: {fname}")
                continue
            # extract responsibility bullets from resume via LLM
            rpt = agent_config['prompts']['resume_exp_extract'] + "\n\nResume:\n" + text
            ind = """Extract the industries where the candidate worked""" + text
            yrs = """Extract the total years of work experience for the candidate""" + text
            desig = """Extract the current designation of the candidate""" + text
            resume_features = llm_call(rpt, agent_config['models']['llm_model'])
            industry = llm_call(ind, agent_config['models']['llm_model'])
            #years_of_exp = llm_call(yrs, agent_config['models']['llm_model'])
            designation = llm_call(desig, agent_config['models']['llm_model'])
            resume_features = clean_text(resume_features)
            industry = clean_text(resume_features)
            #years_of_exp = clean_text(resume_features)
            designation = clean_text(resume_features)
            resumes[fname] = {
                'text': text,
                'resume_features': resume_features,
                'resume_emb': get_embedding(resume_features, embedding_model),
                'industry': industry,
                'ind_emb': get_embedding(industry, embedding_model),
                #'yrs_emb': years_of_exp,
                #'years_of_exp': get_embedding(years_of_exp, embedding_model),
                'designation': designation,
                'desig_emb': get_embedding(designation, embedding_model),
                'full_emb': get_embedding(text, embedding_model),
                'years_extracted': None  # optional future extraction
            }

        # process resumes in resume folder
        for path in folder_files:
            text = load_resume_text_from_path(path)
            if not text:
                logger.warning(f"No text in {path}; skipping")
                continue
            fname = os.path.basename(path)
            rpt = agent_config['prompts']['resume_exp_extract'] + "\n\nResume:\n" + text
            ind = """Extract the industries where the candidate worked""" + text
            yrs = """Extract the total years of work experience for the candidate""" + text
            desig = """Extract the current designation of the candidate""" + text
            resume_features = llm_call(rpt, agent_config['models']['llm_model'])
            industry = llm_call(ind, agent_config['models']['llm_model'])
            #years_of_exp = llm_call(yrs, agent_config['models']['llm_model'])
            designation = llm_call(desig, agent_config['models']['llm_model'])
            resume_features = clean_text(resume_features)
            industry = clean_text(resume_features)
            #years_of_exp = clean_text(resume_features)
            designation = clean_text(resume_features)
            resumes[fname] = {
                'text': text,
                'resume_features': resume_features,
                'resume_emb': get_embedding(resume_features, embedding_model),
                'industry': industry,
                'ind_emb': get_embedding(industry, embedding_model),
                #'yrs_emb': years_of_exp,
                #'years_of_exp': get_embedding(years_of_exp, embedding_model),
                'designation': designation,
                'desig_emb': get_embedding(designation, embedding_model),
                'full_emb': get_embedding(text, embedding_model),
                'years_extracted': None  # optional future extraction
            }
            

        if not resumes:
            return jsonify({'output': 'No valid resumes with extracted content.'}), 400

        logger.info(f"Loaded resumes: {list(resumes.keys())}")

        # compute similarity components for each resume
        weights = agent_config.get('weights', {})
        w_job = float(weights.get('job_features', 0.5))
        w_ind = float(weights.get('industry', 0.15))
        w_role = float(weights.get('role', 0.15))
        w_years = float(weights.get('years', 0.2))

        results = []
        for fname, meta in resumes.items():
            # job_features vs resume_features
            v_job = np.array(job_feat_emb, dtype=float)
            v_resume = np.array(meta['resume_emb'], dtype=float)
            min_len = min(len(v_job), len(v_resume))
            sim_job = 0.0
            if min_len > 0:
                sim_job = float(cosine_similarity([v_job[:min_len]], [v_resume[:min_len]])[0][0])

            # industry similarity: industry_emb vs resume full text embedding
            v_ind = np.array(industry_emb, dtype=float)
            v_res_ind = np.array(meta['ind_emb'], dtype=float)
            min_len2 = min(len(v_ind), len(v_res_ind))
            sim_ind = 0.0
            if min_len2 > 0 and (industry_input):
                sim_ind = float(cosine_similarity([v_ind[:min_len2]], [v_res_ind[:min_len2]])[0][0])

            # role similarity: role_emb vs resume full embedding
            v_role = np.array(role_emb, dtype=float)
            v_res_desig = np.array(meta['desig_emb'], dtype=float)
            min_len3 = min(len(v_role), len(v_res_desig))
            sim_role = 0.0
            if min_len3 > 0 and (role_name):
                sim_role = float(cosine_similarity([v_role[:min_len3]], [v_res_desig[:min_len3]])[0][0])

            # Years similarity: role_emb vs resume full embedding
            #v_yrs = np.array(years_emb, dtype=float)
            #v_res_yrs = np.array(meta['yrs_emb'], dtype=float)
            #min_len3 = min(len(v_role), len(v_res_desig))
            #sim_role = 0.0
            #if min_len3 > 0 and (role_name):
             #   sim_yrs = float(cosine_similarity([v_yrs[:min_len3]], [v_res_yrs[:min_len3]])[0][0])

            #years match: if target_years provided attempt to extract years from resume quick heuristic
            sim_years = 0.0
            if target_years is not None:
                # simple heuristic: search for "Total Experience" or "Experience: X years" in resume text
                txt = meta['text']
                m = re.search(r'(\d{1,2})\s*\+?\s*(?:years|yrs)\b', txt, re.IGNORECASE)
                if m:
                    try:
                        yrs = float(m.group(1))
                        diff = abs(yrs - target_years)
                        # convert difference to score: exact match ->1, difference >=10 ->0
                        sim_years = max(0.0, 1.0 - (diff / 10.0))
                    except Exception:
                        sim_years = 0.0
                else:
                    sim_years = 0.0

            # combine weighted score
            final_score = (w_job * sim_job) + (w_ind * sim_ind) + (w_role * sim_role) + (w_years * sim_years)
            results.append({
                'filename': fname,
                'sim_job': sim_job,
                'sim_industry': sim_ind,
                'sim_role': sim_role,
                'sim_years': sim_years,
                'score': final_score,
                'resume_features': meta.get('resume_features','')[:2000],
                'text': meta.get('text')
            })
            logger.debug(f"{fname} => job:{sim_job:.4f} ind:{sim_ind:.4f} role:{sim_role:.4f} yrs:{sim_years:.4f} final:{final_score:.4f}")

        # sort top N (10)
        top_n = sorted(results, key=lambda x: x['score'], reverse=True)[:10]

        # build candidates_list response
        candidates_list = []
        for idx, r in enumerate(top_n, 1):
            # create fit summary using LLM (optional)
            fit_prompt = agent_config['prompts']['refinement'].format(job_desc=jd_text, resume_summary=r['text'])
            fit_summary = llm_call(fit_prompt, agent_config['models']['llm_model'])
            candidates_list.append({
                'filename': r['filename'],
                'rank': idx,
                'score': float(r['score']),
                'sim_job': float(r['sim_job']),
                'sim_industry': float(r['sim_industry']),
                'sim_role': float(r['sim_role']),
                'sim_years': float(r['sim_years']),
                'fit_summary': fit_summary
            })

        logger.info(f"Returning {len(candidates_list)} candidates")
        logger.info(f"{fit_prompt}")
        return jsonify({'candidates': candidates_list}), 200

    except Exception as e:
        logger.exception("Processing error")
        return jsonify({'output': f"Server error: {e}"}), 500

@app.route('/draft_email', methods=['POST'])
def draft_email_route():
    try:
        data = request.get_json() or {}
        recipient_name = data.get('recipient_name', 'Team')
        sender_signature = data.get('sender_signature', '')
        selected = data.get('selected', [])
        job_desc = data.get('job_desc', '')

        lines = [f"Dear {recipient_name},", "", "Please find shortlisted candidates below:", ""]
        for c in selected:
            fname = c.get('filename') or 'Candidate'
            lines.append(f"- {fname}: Score {c.get('score',0):.3f}. Summary: {c.get('fit_summary','')}")
        lines.append("")
        lines.append("Regards,")
        lines.append(sender_signature or "")
        email_body = "\n".join(lines)
        return jsonify({'email': email_body}), 200
    except Exception as e:
        logger.exception("draft_email error")
        return jsonify({'output': f"Error: {e}"}), 500

@app.route('/generate_questions', methods=['POST'])
def generate_questions_route():
    try:
        data = request.get_json() or {}
        job_desc = data.get('job_desc', '')
        candidate_profile = data.get('candidate_profile', '')
        if not job_desc or not candidate_profile:
            return jsonify({'output': 'Provide job_desc and candidate_profile'}), 400

        # simple placeholder questions; you can call LLM here
        questions = [
            "1. Please walk me through your current role and responsibilities.",
            "2. What relevant technologies and tools have you used for this role?",
            "3. Explain a recent project where you used the core skills mentioned in the JD.",
            "4. Why are you looking to leave your current employer?",
            "5. What are your current and expected CTC?",
            "6. Are you open to relocation or remote work?",
            "7. Describe a challenging situation at work and how you resolved it.",
            "8. How do you prioritize tasks when handling multiple projects?",
            "9. How do you handle conflicts within a team?",
            "10. What are your notice period and earliest joining date?"
        ]
        return jsonify({'questions': '\n'.join(questions)}), 200
    except Exception as e:
        logger.exception("generate_questions error")
        return jsonify({'output': f"Error: {e}"}), 500

# HTML template is expected to be provided separately; keep placeholder
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8"/>
  <meta name="viewport" content="width=device-width,initial-scale=1"/>
  <title>HR Agent</title>
  <script src="https://cdn.tailwindcss.com"></script>
</head>
<body class="bg-gray-100 min-h-screen p-6">
  <div class="max-w-5xl mx-auto bg-white rounded shadow p-6">
    <h1 class="text-2xl font-bold mb-4">HR Resume Agent</h1>
    <div>
      <nav class="mb-4">
        <button id="tab-analysis" class="px-4 py-2 bg-blue-600 text-white rounded mr-2">Resume Analysis</button>
        <button id="tab-email" class="px-4 py-2 bg-gray-200 rounded mr-2">Draft Email</button>
        <button id="tab-questions" class="px-4 py-2 bg-gray-200 rounded">Interview Questions</button>
      </nav>

      <div id="panel-analysis">
    <h2 class="text-lg font-semibold mb-2">Resume Analysis</h2>

    <label class="block mb-2">Job Description (paste):</label>
    <textarea id="job_desc" rows="4" class="w-full border rounded p-2 mb-3" placeholder="Enter job description"></textarea>

    <label class="block mb-2">Upload Job Description (optional .doc/.docx/.pdf):</label>
    <input id="jd_file" type="file" accept=".doc,.docx,.pdf" class="mb-3"/>

    <div class="grid grid-cols-1 md:grid-cols-3 gap-3 mb-3">
      <div>
        <label class="block mb-1">Years of Experience (target)</label>
        <input id="years_exp" class="w-full border rounded p-2" placeholder="e.g., 5 or 3-7"/>
      </div>
      <div>
        <label class="block mb-1">Industry</label>
        <input id="industry" class="w-full border rounded p-2" placeholder="e.g., Manufacturing"/>
      </div>
      <div>
        <label class="block mb-1">Role</label>
        <select id="role_name" class="w-full border rounded p-2">
          <option value="">Not Applicable</option>
          <option>CEO</option>
          <option>CFO</option>
          <option>Vice President</option>
          <option>General Manager</option>
          <option>Senior Manager</option>
          <option>Manager</option>
          <option>Area Manager</option>
          <option>Not Applicable</option>
        </select>
      </div>
    </div>

    <label class="block mb-2">Upload Resumes (.doc, .docx, .pdf):</label>
    <input id="file_input" type="file" accept=".doc,.docx,.pdf,application/msword,application/vnd.openxmlformats-officedocument.wordprocessingml.document" multiple class="mb-3"/>
    <button id="analyze_btn" class="px-4 py-2 bg-blue-600 text-white rounded">Analyze & Find Top 5</button>

    <div class="mt-4">
      <label class="font-medium">Matched Candidates (select to include in email):</label>
      <div id="matches" class="mt-2"></div>
    </div>
  </div>

      <div id="panel-email" style="display:none;">
        <h2 class="text-lg font-semibold mb-2">Draft Email</h2>
        <label class="block mb-2">Recipient Name:</label>
        <input id="recipient_name" class="w-full border rounded p-2 mb-3"/>
        <label class="block mb-2">Sender Signature (name & title):</label>
        <input id="sender_signature" class="w-full border rounded p-2 mb-3"/>
        <button id="draft_email_btn" class="px-4 py-2 bg-green-600 text-white rounded mb-3">Draft Email for Selected Candidates</button>
        <label class="block mb-2">Email Draft:</label>
        <textarea id="email_output" rows="10" class="w-full border rounded p-2" readonly></textarea>
      </div>

      <div id="panel-questions" style="display:none;">
        <h2 class="text-lg font-semibold mb-2">Interview Questions</h2>
        <label class="block mb-2">Job Description (paste or upload):</label>
        <textarea id="job_desc_q" rows="4" class="w-full border rounded p-2 mb-3" placeholder="Enter job description"></textarea>
        <label class="block mb-2">Candidate Profile (paste or upload):</label>
        <textarea id="candidate_profile_q" rows="4" class="w-full border rounded p-2 mb-3" placeholder="Paste candidate summary/profile"></textarea>
        <button id="gen_questions_btn" class="px-4 py-2 bg-indigo-600 text-white rounded mb-3">Generate Questions</button>
        <label class="block mb-2">Interview Questions:</label>
        <textarea id="questions_output" rows="12" class="w-full border rounded p-2" readonly></textarea>
      </div>
    </div>
  </div>

<script>
  // Tab handling
  const tabAnalysis = document.getElementById('tab-analysis');
  const tabEmail = document.getElementById('tab-email');
  const tabQuestions = document.getElementById('tab-questions');
  const panelAnalysis = document.getElementById('panel-analysis');
  const panelEmail = document.getElementById('panel-email');
  const panelQuestions = document.getElementById('panel-questions');

  function showPanel(panel) {
    panelAnalysis.style.display = 'none';
    panelEmail.style.display = 'none';
    panelQuestions.style.display = 'none';
    tabAnalysis.className = 'px-4 py-2 bg-gray-200 rounded mr-2';
    tabEmail.className = 'px-4 py-2 bg-gray-200 rounded mr-2';
    tabQuestions.className = 'px-4 py-2 bg-gray-200 rounded';
    if (panel === 'analysis') {
      panelAnalysis.style.display = 'block';
      tabAnalysis.className = 'px-4 py-2 bg-blue-600 text-white rounded mr-2';
    } else if (panel === 'email') {
      panelEmail.style.display = 'block';
      tabEmail.className = 'px-4 py-2 bg-blue-600 text-white rounded mr-2';
    } else {
      panelQuestions.style.display = 'block';
      tabQuestions.className = 'px-4 py-2 bg-blue-600 text-white rounded';
    }
  }

  tabAnalysis.onclick = () => showPanel('analysis');
  tabEmail.onclick = () => showPanel('email');
  tabQuestions.onclick = () => showPanel('questions');

  // Analysis logic
  document.getElementById('analyze_btn').addEventListener('click', async () => {
    const jobDesc = document.getElementById('job_desc').value;
    const files = document.getElementById('file_input').files;
    if (!jobDesc) { alert('Please enter job description'); return; }
    const fd = new FormData();
    fd.append('text_input', jobDesc);
    for (let i=0;i<files.length;i++) fd.append('file_input', files[i]);
    document.getElementById('matches').innerHTML = 'Processing...';
    const resp = await fetch('/process', { method: 'POST', body: fd });
    const data = await resp.json();
    if (!resp.ok) {
      document.getElementById('matches').innerText = data.output || 'Error';
      return;
    }
    // data.output is a structured JSON in our implementation (list of candidates)
    const candidates = data.candidates || [];
    if (candidates.length === 0) {
      document.getElementById('matches').innerText = 'No matches found.';
      return;
    }
    // render checkboxes
    const container = document.getElementById('matches');
    container.innerHTML = '';
    candidates.forEach((c, idx) => {
      const div = document.createElement('div');
      div.className = 'flex items-start mb-2';
      div.innerHTML = `<input type="checkbox" data-idx="${idx}" id="cand_${idx}" class="mr-2 mt-1"><div><b>${c.filename}</b> (Score: ${c.score.toFixed(2)})<div class="text-sm text-gray-700">${c.fit_summary}</div></div>`;
      container.appendChild(div);
    });
    // store candidates in window for email drafting
    window._candidates = candidates;
  });

  // Draft email
  document.getElementById('draft_email_btn').addEventListener('click', async () => {
    const recipient = document.getElementById('recipient_name').value || 'Team';
    const signature = document.getElementById('sender_signature').value || '';
    const candidates = window._candidates || [];
    const selected = [];
    document.querySelectorAll('#matches input[type=checkbox]').forEach(cb => {
      if (cb.checked) {
        const idx = parseInt(cb.getAttribute('data-idx'));
        if (!isNaN(idx) && candidates[idx]) selected.push(candidates[idx]);
      }
    });
    if (selected.length === 0) { alert('Select at least one candidate'); return; }
    document.getElementById('email_output').value = 'Drafting email...';
    const resp = await fetch('/draft_email', {
      method: 'POST',
      headers: {'Content-Type':'application/json'},
      body: JSON.stringify({recipient_name: recipient, sender_signature: signature, selected: selected, job_desc: document.getElementById('job_desc').value})
    });
    const data = await resp.json();
    if (resp.ok) {
      document.getElementById('email_output').value = data.email;
      showPanel('email');
    } else {
      document.getElementById('email_output').value = data.output || 'Error drafting email';
    }
  });

  // Generate questions
  document.getElementById('gen_questions_btn').addEventListener('click', async () => {
    const jd = document.getElementById('job_desc_q').value;
    const profile = document.getElementById('candidate_profile_q').value;
    if (!jd || !profile) { alert('Provide job description and candidate profile'); return; }
    document.getElementById('questions_output').value = 'Generating questions...';
    const resp = await fetch('/generate_questions', {
      method:'POST',
      headers:{'Content-Type':'application/json'},
      body: JSON.stringify({job_desc: jd, candidate_profile: profile})
    });
    const data = await resp.json();
    if (resp.ok) document.getElementById('questions_output').value = data.questions;
    else document.getElementById('questions_output').value = data.output || 'Error';
    showPanel('questions');
  });

</script>
</body>
</html>"""

def open_browser(port=8080, delay=1.5):
    try:
        time.sleep(delay)
        webbrowser.open_new(f'http://127.0.0.1:{port}/')
    except Exception as e:
        logger.error(f"Browser open failed: {e}")

if __name__ == '__main__':
    threading.Thread(target=open_browser, args=(8080,1.5), daemon=True).start()
    app.run(host='127.0.0.1', port=8080, debug=False)


# In[ ]:




